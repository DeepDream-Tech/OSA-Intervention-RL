from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.left_margin   = Inches(1.18)
section.right_margin  = Inches(1.18)
section.top_margin    = Inches(1.18)
section.bottom_margin = Inches(1.18)

def set_font(run, name_zh="宋体", name_en="Times New Roman",
             size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_zh)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "黑体", "Arial", size=18, bold=True, color=(0x1F, 0x49, 0x7D))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(12)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "黑体", "Arial", size=14, bold=True, color=(0x2E, 0x74, 0xB5))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "黑体", "Arial", size=12, bold=True, color=(0x1F, 0x49, 0x7D))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_font(run, "宋体", "Times New Roman", size=11)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    set_font(run, "宋体", "Times New Roman", size=11)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(17)

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    run = p.add_run(text)
    set_font(run, "宋体", "Times New Roman", size=10, color=(0x59, 0x59, 0x59))
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ═══════════════ 正文 ═══════════════

add_title(doc, "OSA 声学干预强化学习方案工程评估报告")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("针对 INTERVENTION_RL_MODEL_SPEC v1.1 的工程角度问题分析")
set_font(run, "仿宋", "Times New Roman", size=11, color=(0x70, 0x70, 0x70))
p.paragraph_format.space_after = Pt(16)

add_note(doc, (
    "本报告从工程可行性与逻辑一致性两个维度，对上述 RL 建模规范进行系统评估，"
    "分为【核心架构问题】【可行性问题】【逻辑问题】【工程实现风险】四大类，共 11 项具体问题。"
))

# ── 一、核心架构问题 ──────────────────────────────────────
add_h1(doc, "一、核心架构问题")

add_h2(doc, "问题 1：动作空间过度简化与目标期望之间存在根本矛盾")
add_body(doc, (
    "方案将可学习动作限定为响度（6 个离散档位），波形、频率、时长全部固定。这是合理的第一步，"
    "但方案同时期望 RL 策略能在 60 秒内通过响度变化让生理指标恢复正常——二者之间存在根本矛盾："
))
add_bullet(doc, "OSA 的生理恢复本身需要数十秒，而整个 episode 只有 60 秒。")
add_bullet(doc, (
    "策略最多执行 12 次动作（N_max = T_active_max / delta = 12），"
    "信号延迟 + 生理响应延迟 + 恢复资格门控（T_min_active = 10 s）会进一步压缩有效步数。"
))
add_bullet(doc, "实际可用于学习的有效交互步数极少，RL 在此类稀疏 episode 上很难收敛。")

add_h2(doc, "问题 2：奖励信号设计存在严重的反馈延迟")
add_body(doc, "奖励的主要正向分量是 delta_H_k（综合生理负担变化），但：")
add_bullet(doc, "PPG 特征窗口为 15 秒，心率变化要延迟至少一个特征窗口才反映到 H_t。")
add_bullet(doc, "声音播放期间 m_k^A = 0，Audio 特征被屏蔽，进一步减少可用信息。")
add_bullet(doc, "信用分配只回溯 3 步（15 秒），但 OSA 的生理响应延迟可能远超此范围。")
add_bullet(doc, "结果：策略几乎无法从 delta_H_k 中区分哪个响度档位起了作用。")

add_h2(doc, "问题 3：关键代理奖励黑盒化（Proxy Reward Hacking）")
add_body(doc, (
    "方案用综合生理负担 H_t 作为代理奖励，各分量权重 w = [0.35, 0.30, 0.25, 0.10] "
    "来源是人工设定，且大量关键参数标注为 [TBD-CAL]。RL 会优化这个代理，而非真实的 OSA 预防效果。"
))
add_note(doc, (
    "方案自述：如果没有真实事件标签，实验结论只能表述为「降低风险/改善代理指标」——"
    "但奖励主体 w_H * delta_H_k 本身就是代理，策略完全可能学到让 H_t 数值下降而非真正预防 OSA 的行为。"
))

# ── 二、可行性问题 ────────────────────────────────────────
add_h1(doc, "二、可行性问题")

add_h2(doc, "问题 4：离线 RL 的数据需求与实际可收集数据之间存在矛盾")
add_body(doc, "方案推荐使用 CQL / IQL 进行离线强化学习，但存在以下障碍：")
add_bullet(doc, (
    "基线策略（递增响度）为确定性行为，不覆盖响度 = 0 的静默观察动作，"
    "离线数据的动作分布极度集中，CQL 的支持外推问题将非常严重。"
))
add_bullet(doc, (
    "OSA 事件本身稀疏（UCDDB 数据中 Apnea 仅占 3.4%），真实触发 episode 数量极少，"
    "不足以支撑离线 RL 的样本需求。"
))

add_h2(doc, "问题 5：恢复资格门控设计与预测性触发逻辑相互矛盾")
add_body(doc, (
    "恢复条件 E_k 的第二路径要求 H_k^peak >= theta_enter（0.45）才能进入恢复流程。"
    "若触发模型提前发出预警，触发时生理指标尚未恶化，H_t 可能从未超过 0.45，"
    "导致整个 episode 只能以超时或人工停止结束，策略永远无法获得 B_rec = 8.0 的成功奖励。"
    "这与方案第 1 节「外部模型预测未来事件」的设计前提直接冲突。"
))

add_h2(doc, "问题 6：空格键作为觉醒标签的实验设计缺陷")
add_body(doc, "方案以空格键作为唯一人工觉醒确认，存在系统性的假阴性问题：")
add_bullet(doc, "真实睡眠实验中，受试者在睡着时不会按键，只有醒来后才能操作。")
add_bullet(doc, "策略会被惩罚 C_awake = 8.0，但没有按键不等于没有觉醒。")
add_bullet(doc, (
    "大量真实觉醒但未按键的情况，会被错误记录为无惩罚的「恢复成功」，严重污染训练信号。"
))

# ── 三、逻辑问题 ──────────────────────────────────────────
add_h1(doc, "三、逻辑问题")

add_h2(doc, "问题 7：奖励函数各项之间存在互相冲突的激励")
add_body(doc, (
    "方案同时引入以下惩罚项：lambda_L（抑制高响度）、lambda_D（抑制声音剂量）、"
    "lambda_t（每步扣分鼓励尽快结束）、C_timeout（超时惩罚）。"
    "如果响度不够高无法触发恢复，策略超时被扣 4.0，"
    "但若提高响度又受 lambda_L、lambda_D 的惩罚。"
))
add_body(doc, (
    "这几个项之间缺乏明确的量纲对齐，权重比例完全依赖经验。"
    "实际训练中策略极可能陷入局部最优：保持低响度等待超时，"
    "比尝试高响度触发恢复的期望代价更低。"
))

add_h2(doc, "问题 8：ARMED -> INTERVENING 转换要求三个模态同时可用")
add_body(doc, (
    "就绪条件为 C_t^ready = 1[b^base=1] x m_t^A x m_t^P x m_t^I，"
    "三个模态需同时满足质量阈值。"
))
add_body(doc, (
    "真实睡眠场景中 PPG 信号因运动伪影或接触问题频繁失效，"
    "将导致系统大量无法进入干预状态。T_arm = 10 s 等待超时后直接进入 SAFE_STOP，"
    "实际上会使大量本可干预的 episode 失效。"
))

add_h2(doc, "问题 9：个体基线与人群基线的数据泄漏风险")
add_body(doc, (
    "方案要求按受试者划分训练 / 验证 / 测试集，"
    "但若人群基线来自训练集，测试集受试者使用该人群基线时会引入隐含的分布偏移。"
    "此外，T_base = 120 s 的基线窗口要求 episode 开始前已有稳定数据，"
    "在真实部署场景中难以保证。"
))

# ── 四、工程实现风险 ──────────────────────────────────────
add_h1(doc, "四、工程实现风险")

add_h2(doc, "问题 10：实时性要求未进行预算评估")
add_body(doc, "方案要求：")
add_bullet(doc, "100 ms 内响应空格键（手动静音延迟约束）。")
add_bullet(doc, (
    "每 1 秒内完成 Audio（8000 个样本）、PPG（25 个样本）、IMU（33 个样本）"
    "的特征提取、融合与策略推理。"
))
add_body(doc, (
    "上述计算负载在嵌入式 / 移动端设备上的可行性完全未作分析，存在较大落地风险。"
))

add_h2(doc, "问题 11：声学回声消除是核心假设但缺乏实现路径")
add_body(doc, (
    "方案多处依赖回声消除来防止播放声音污染麦克风特征，"
    "但第 1.1 节仅建议「使用播放参考信号进行 AEC」，"
    "未给出具体算法选择、延迟估计或耳机场景下的可行性分析。"
))
add_note(doc, (
    "耳机扬声器到麦克风的直达声路径极短、时延极小，传统 AEC 在此场景下效果不稳定；"
    "若 AEC 失败，播放声音将直接污染 Audio 特征，破坏整个恢复判定链路。"
))

# ── 五、总结与建议 ────────────────────────────────────────
add_h1(doc, "五、总结与建议")

add_body(doc, (
    "该方案在形式上是完整的——符号定义、约束、日志、验收条件均有涉及，"
    "体现了对 RL 建模的基本理解。但从工程可行性角度看，"
    "核心问题是：RL 在该场景下是否真的能学到有用策略？"
))

add_body(doc, "主要制约因素：")
add_bullet(doc, "Episode 极短（60 s），有效交互步数不超过 12 步。")
add_bullet(doc, "数据稀疏（OSA 触发率约 3.4%），离线数据量严重不足。")
add_bullet(doc, "奖励延迟高（PPG 窗口 15 s），信用分配困难。")
add_bullet(doc, "代理奖励质量不确定，可能导致策略优化方向偏离真实目标。")

add_body(doc, "建议的更务实路径：")
add_bullet(doc, (
    "优先用基线策略（递增响度）运行足够数量的真实 episode，"
    "验证 H_t 的变化是否与响度存在可预测的相关性。"
))
add_bullet(doc, (
    "在数据充足、代理奖励经过验证之后，再引入离线 RL 组件；"
    "初期应保持 Shadow Mode 对比，不直接上线 RL 策略。"
))
add_bullet(doc, (
    "完成 [TBD-CAL] 标注的所有声学标定参数，"
    "并由临床和听力安全负责人审核 L_min、L_max、D_max 之后，"
    "才具备进行真实受试者实验的前提条件。"
))

out_path = "/Users/liujia/Desktop/OSA-Intervention-RL/OSA_RL_方案工程评估报告.docx"
doc.save(out_path)
print("Saved:", out_path)
